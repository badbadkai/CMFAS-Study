"""Parse the HI and RES5 mock papers into their mocks.json files.

These papers use a "Qn" marker layout (distinct from the M9/M9A "12." layout):
  - A paragraph "Q1" marks a question; the stem is the paragraph(s) that
    follow, up to the first option line.
  - Options are lines "A. ..." either newline-split in one paragraph (HI)
    or one paragraph per option (RES5).
  - Answers live in Q/Answer tables (5 Q/Ans column pairs per row).

RES5 has two parts: Part 1 (Q1-110) and Part 2 (Q1-40). Part 2 restarts its
numbering, so its questions and answers are offset by 110 to stay unique
within the paper (Part 2 Q1 -> 111 ... Q40 -> 150).
"""
import json
import re
from pathlib import Path

import docx

SRC = Path.home() / "Downloads"
ROOT = Path(__file__).resolve().parent.parent / "src" / "data"

QMARK_RE = re.compile(r"^Q(\d+)$")
OPT_RE = re.compile(r"^([ABCD])[.)]\s*(.*)$")

CONFIGS = {
    "HI": {
        "files": [f"CMFAS_HI_Mock_Examination_Paper_{n}.docx" for n in range(1, 5)],
        "out": ROOT / "hi" / "mocks.json",
    },
    "RES5": {
        "files": [f"RES5_Paper_{n}.docx" for n in range(1, 5)],
        "out": ROOT / "res5" / "mocks.json",
    },
}


def answers_from_tables(doc):
    """Read every Q/Ans table. Table index i offsets its question numbers by i*110
    so a two-part paper (RES5) keeps Part 2's restarted numbering unique."""
    answers = {}
    for ti, tb in enumerate(doc.tables):
        header = [c.text.strip().lower() for c in tb.rows[0].cells]
        if not any(h.startswith("ans") for h in header):
            continue
        offset = ti * 110
        for row in tb.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            for i in range(0, len(cells) - 1, 2):
                q, a = cells[i], cells[i + 1].upper()
                if q.isdigit() and a in ("A", "B", "C", "D"):
                    answers[int(q) + offset] = a
    return answers


def parse_paper(path: Path, code: str, paper_num: int):
    doc = docx.Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    answers = answers_from_tables(doc)

    questions = []
    offset = 0
    i = 0
    while i < len(paras):
        t = paras[i]
        if t.upper().startswith("PART 2"):
            offset = 110
            i += 1
            continue
        m = QMARK_RE.match(t)
        if not m:
            i += 1
            continue
        num = int(m.group(1)) + offset

        # Gather stem paragraphs until the first option line.
        j = i + 1
        stem_parts = []
        while j < len(paras):
            if QMARK_RE.match(paras[j]) or paras[j].upper().startswith("PART 2"):
                break
            if any(OPT_RE.match(ln.strip()) for ln in paras[j].split("\n")):
                break
            stem_parts.append(paras[j])
            j += 1
        stem = " ".join(stem_parts).strip()

        # Collect options until the next question / part / four found.
        opts = {}
        while j < len(paras) and len(opts) < 4:
            if QMARK_RE.match(paras[j]) or paras[j].upper().startswith("PART 2"):
                break
            for ln in (x.strip() for x in paras[j].split("\n") if x.strip()):
                om = OPT_RE.match(ln)
                if om and om.group(1) not in opts:
                    opts[om.group(1)] = om.group(2).strip()
            j += 1

        if stem and {"A", "B", "C", "D"} <= set(opts):
            ans = answers.get(num)
            if ans is None:
                raise SystemExit(f"{code} Paper {paper_num} Q{num}: no answer found")
            questions.append(
                {
                    "id": f"{code}-P{paper_num}-Q{num}",
                    "paper": paper_num,
                    "num": num,
                    "stem": stem,
                    "options": {k: opts[k] for k in ("A", "B", "C", "D")},
                    "answer": ans,
                }
            )
        i = j
    return questions


def main():
    for code, cfg in CONFIGS.items():
        papers = []
        for n, fname in enumerate(cfg["files"], start=1):
            qs = parse_paper(SRC / fname, code, n)
            papers.append({"paper": n, "questions": qs})
            print(f"{code} Paper {n}: {len(qs)} questions")
        cfg["out"].parent.mkdir(parents=True, exist_ok=True)
        cfg["out"].write_text(
            json.dumps({"module": code, "papers": papers}, indent=1, ensure_ascii=False),
            encoding="utf-8",
        )
        print(f"  wrote {sum(len(p['questions']) for p in papers)} questions to {cfg['out']}\n")


if __name__ == "__main__":
    main()
