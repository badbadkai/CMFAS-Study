"""Parse the 4 CMFAS_M9A_Mock_Paper_*.docx into src/data/m9a/mocks.json.

The four papers use slightly different layouts, so the parser is tolerant:
  - Stems appear as paragraphs like "12. stem" or "12: stem".
  - Options follow the stem either as one newline-separated block
    ("A. ...\nB. ...") or as one paragraph per option ("A. ...").
  - Answers come from either a flat list after an "ANSWER KEY" heading
    (paper 1) or a Q/Ans table (papers 2-4).
"""
import json
import re
from pathlib import Path

import docx

SRC = Path.home() / "Downloads"
OUT = Path(__file__).resolve().parent.parent / "src" / "data" / "m9a" / "mocks.json"

STEM_RE = re.compile(r"^(\d+)[.:]\s+(.+)$", re.S)
OPT_RE = re.compile(r"^([ABCD])[.)]\s*(.*)$")


def answers_from_table(doc):
    answers = {}
    for tb in doc.tables:
        header = [c.text.strip().lower() for c in tb.rows[0].cells]
        if "ans" not in header:
            continue
        for row in tb.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            for i in range(0, len(cells) - 1, 2):
                q, a = cells[i], cells[i + 1].upper()
                if q.isdigit() and a in ("A", "B", "C", "D"):
                    answers[int(q)] = a
    return answers


def answers_from_flat(paras):
    for i, t in enumerate(paras):
        if t.upper() == "ANSWER KEY":
            tail = [x.upper() for x in paras[i + 1 :] if x.strip().upper() in ("A", "B", "C", "D")]
            return {n + 1: a for n, a in enumerate(tail)}
    return {}


def parse_paper(path: Path, paper_num: int):
    doc = docx.Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    answers = answers_from_flat(paras) or answers_from_table(doc)

    # Cut the answer-key heading off the body so it is not mistaken for content.
    body = paras
    for i, t in enumerate(paras):
        if t.upper() == "ANSWER KEY":
            body = paras[:i]
            break

    questions = []
    i = 0
    while i < len(body):
        m = STEM_RE.match(body[i])
        if not m or OPT_RE.match(body[i]):
            i += 1
            continue
        num = int(m.group(1))
        stem = m.group(2).strip()
        opts = {}
        j = i + 1
        while j < len(body) and len(opts) < 4:
            if STEM_RE.match(body[j]) and not OPT_RE.match(body[j]):
                break
            for ln in (x.strip() for x in body[j].split("\n") if x.strip()):
                om = OPT_RE.match(ln)
                if om and om.group(1) not in opts:
                    opts[om.group(1)] = om.group(2).strip()
            j += 1
        if {"A", "B", "C", "D"} <= set(opts):
            questions.append({"num": num, "stem": stem, "options": opts})
            i = j
        else:
            i += 1

    out = []
    for q in questions:
        ans = answers.get(q["num"])
        if ans is None:
            raise SystemExit(f"Paper {paper_num} Q{q['num']}: no answer found")
        out.append(
            {
                "id": f"M9A-P{paper_num}-Q{q['num']}",
                "paper": paper_num,
                "num": q["num"],
                "stem": q["stem"],
                "options": {k: q["options"][k] for k in ("A", "B", "C", "D")},
                "answer": ans,
            }
        )
    return out


def main():
    papers = []
    for n in range(1, 5):
        qs = parse_paper(SRC / f"CMFAS_M9A_Mock_Paper_{n}.docx", n)
        papers.append({"paper": n, "questions": qs})
        print(f"Paper {n}: {len(qs)} questions")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(
        json.dumps({"module": "M9A", "papers": papers}, indent=1, ensure_ascii=False),
        encoding="utf-8",
    )
    print(f"Wrote {sum(len(p['questions']) for p in papers)} questions to {OUT}")


if __name__ == "__main__":
    main()
